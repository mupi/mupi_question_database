from django.views.generic import DetailView, ListView, RedirectView, UpdateView, TemplateView, CreateView, DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect

from docx import *
from docx.shared import Inches

import json
import datetime
import io

from .models import Question, Question_List

class QuestionDetailView(LoginRequiredMixin, DetailView):
    model = Question
    template_name = "questions/question_detail.html"
    context_object_name = "question"

class QuestionListView(LoginRequiredMixin, ListView):
    model = Question
    template_name = "questions/question_list.html"
    context_object_name = "question_list"
    paginate_by = 10

class Question_ListDeleteView(LoginRequiredMixin, DeleteView):
    model = Question_List
    template_name = "questions/question_list_detail.html"
    context_object_name = "question_list"
    success_url = "/questions/question_lists/"

class Question_ListCreateView(LoginRequiredMixin, CreateView):
    model = Question_List
    fields = ['question_list_header']
    template_name = "questions/question_selectedList.html"

    def get_context_data(self, *args, **kwargs):
        context = super(CreateQuestionListView, self).get_context_data(**kwargs)
        if not 'checked_questions' in self.request.session or not self.request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = self.request.session['checked_questions']

        context['checked_questions'] = Question.objects.filter(pk__in=checked_questions)
        return context

    def form_valid(self, form):
        new_list = form.save(commit=False)
        new_list.owner = self.request.user

        if not 'checked_questions' in self.request.session or not self.request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = self.request.session['checked_questions']
        new_list.save()

        questions = Question.objects.filter(pk__in=checked_questions)
        for question in questions:
            print(question)
            new_list.questions.add(question)
        new_list.save()

        self.request.session['checked_questions'] = []
        return HttpResponseRedirect("/questions/question_lists/" + str(new_list.pk) + "/")

def check_question(request):
    if (request.method == 'POST'):
        questionPk = request.POST.get('questionPK')
        checked = request.POST.get('checked')

        if not 'checked_questions' in request.session or not request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = request.session['checked_questions']

        if checked == 'true':
            checked_questions.append(int(questionPk))
        else:
            checked_questions = [item for item in checked_questions if item != int(questionPk)]

        request.session['checked_questions'] = checked_questions

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps({"nothing to see": "this isn't happening"}),
            content_type="application/json",
            status = 400
        )

def clear_questions(request):
    if (request.method == 'GET'):

        request.session['checked_questions'] = []

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps({"nothing to see": "this isn't happening"}),
            content_type="application/json",
            status = 400
        )

def list_generator(request):
    if not 'checked_questions' in request.session or not request.session['checked_questions']:
        print('No questions selected')
        # no checked questions
    checked_questions = request.session['checked_questions']
    checked_questions = Question.objects.filter(pk__in=checked_questions)

    document = Document()
    docx_title="Test_List.docx"
    document.add_paragraph(
        "Lista gerada em {:s} as {:s}.".format(
            datetime.date.today().strftime('%d/%m/%Y'),
            datetime.datetime.today().strftime('%X')
        )
    )

    questionCounter = 1
    for question in checked_questions:
        document.add_heading('Question ' + str(questionCounter), level=2)
        p = document.add_paragraph('(')
        p.add_run(question.question_header).bold = True
        p.add_run(')'+question.question_text)
        itemChar = 'a'
        for answer in question.answer_set.all():
            document.add_paragraph(itemChar + ') ' + answer.answer_text)
            itemChar = chr(ord(itemChar) + 1)
        questionCounter = questionCounter + 1
        p = document.add_paragraph()

    document.add_page_break()

    document.save(docx_title)
    data = open(docx_title, "rb").read()

    response = HttpResponse(
        data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    return response

class Question_ListListView(LoginRequiredMixin, ListView):
    model = Question_List
    template_name = "questions/question_list_list.html"
    context_object_name = "question_list"
    success_url = "/questions"
